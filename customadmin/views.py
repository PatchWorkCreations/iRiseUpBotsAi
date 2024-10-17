from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse
from myapp.forms import CourseForm, SubCourseForm, LessonForm
from myapp.models import Course, SubCourse, Lesson
from django.template.loader import render_to_string
from django.http import JsonResponse
import json
import os
from django.conf import settings
from django.contrib.auth.views import LoginView, LogoutView
from django.urls import reverse_lazy


def dashboard(request):
    courses = Course.objects.all()
    return render(request, 'customadmin/dashboard.html', {'courses': courses})

# views.py
from django.shortcuts import redirect, get_object_or_404
from myapp.models import Course

def toggle_course_status(request, pk):
    course = get_object_or_404(Course, pk=pk)
    if request.method == 'POST':
        course.is_active = not course.is_active
        course.save()
    return redirect('dashboard')  # Redirect to the dashboard


def course_list(request):
    courses = Course.objects.all()
    return render(request, 'customadmin/course_list.html', {'courses': courses})

def course_detail(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    sub_courses = course.sub_courses.all()
    return render(request, 'myapp/course_detail.html', {'course': course, 'sub_courses': sub_courses})

from docx import Document
import json
from myapp.models import SubCourse, Lesson

def extract_subcourses_lessons_from_docx(docx_file, course):
    doc = Document(docx_file)

    current_subcourse = None
    current_lesson = None
    lesson_blocks = []  # To hold all the blocks (headers, paragraphs, questions) for a lesson
    block_content = ""  # To accumulate paragraph content
    subcourse_order = SubCourse.objects.filter(parent_course=course).count()  # Set initial order based on existing subcourses
    is_question_block = False  # Track if we are in a question block (e.g., multiple choice)
    current_question = None
    multiple_choices = []
    correct_answer = None  # To track the correct answer for multiple choice questions

    # Function to save the multiple-choice question block
    def save_current_question():
        nonlocal current_question, multiple_choices, lesson_blocks, correct_answer
        if current_question and multiple_choices:
            lesson_blocks.append({
                'type': 'multiple_choice',
                'question': current_question,
                'options': multiple_choices,
                'correct_answer': correct_answer
            })
            current_question = None
            multiple_choices = []
            correct_answer = None

    # Function to detect reflection questions (starting with a number)
    def is_reflection_question(text):
        return text and text[0].isdigit()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check for multiple-choice or short-answer question types
        is_bold = any(run.bold for run in para.runs if run.bold)
        is_italic = any(run.italic for run in para.runs if run.italic)  # Check for italic text (for answer and questions)
        is_short_header = len(text) <= 50 and "\n" not in text
        starts_with_bullet = text.startswith('•') or text.startswith('-')  # Bullet points or dashes

        # Detect Multiple-Choice Questions block
        if is_italic and not is_question_block:
            is_question_block = True

        # Handle multiple-choice questions
        if is_question_block:
            # Start a new question
            if text.startswith(('What', 'Which', 'How', 'Why', 'Who', 'Where', 'When', 'If', 'Do')):
                save_current_question()  # Save the previous question before starting a new one
                current_question = text
                continue

            # Add options (A., B., C., etc.)
            elif text.startswith(('A.', 'B.', 'C.', 'D.')):
                multiple_choices.append(text)
                continue

            # Handle the correct answer separately (bold + italicized text)
            elif text.startswith("Answer:") or is_italic:
                correct_answer = text.split('Answer:')[-1].strip()  # Store the correct answer separately
                save_current_question()  # Save the multiple-choice block
                is_question_block = False  # Exit the multiple-choice block
                continue

        # Handle reflection questions starting with a number
        if is_reflection_question(text):
            lesson_blocks.append({
                'type': 'question',
                'content': text  # Save as a question block
            })
            continue

        # Indent and format bullet points
        if starts_with_bullet:
            block_content = f"<div style='margin-left: 20px;'>{text}</div>"
            lesson_blocks.append({'type': 'paragraph', 'content': block_content})
            continue

        # Handle Subcourses
        if "Sub course" in text and ':' in text:
            # Save the previous lesson's content
            if current_lesson and lesson_blocks:
                current_lesson.content = json.dumps(lesson_blocks)
                current_lesson.save()
                lesson_blocks = []

            # Create or get subcourse
            subcourse_title = text.split(':', 1)[1].strip()
            current_subcourse, created = SubCourse.objects.get_or_create(
                title=subcourse_title, parent_course=course,
                defaults={'order': subcourse_order + 1, 'units': 1, 'hours': 1}
            )
            subcourse_order += created

            # Reset lesson order
            lesson_order = Lesson.objects.filter(parent_sub_course=current_subcourse).count()
            current_lesson = None

        # Handle Lessons
        elif "Lesson" in text and ':' in text and current_subcourse:
            # Save the previous lesson's content
            if current_lesson and lesson_blocks:
                current_lesson.content = json.dumps(lesson_blocks)
                current_lesson.save()

            # Create or get lesson
            lesson_title = text.split(':', 1)[1].strip()
            current_lesson, created = Lesson.objects.get_or_create(
                title=lesson_title, parent_sub_course=current_subcourse,
                defaults={'order': lesson_order + 1}
            )
            lesson_order += created
            lesson_blocks = []
            inside_lesson = True

        # Handle Headers (e.g., "Key Points", "Exercises")
        elif (is_bold or is_short_header) and current_lesson and inside_lesson:
            lesson_blocks.append({'type': 'header', 'content': text})

        # Handle Paragraphs
        elif current_lesson and inside_lesson:
            block_content += text + " "
            if block_content:
                lesson_blocks.append({'type': 'paragraph', 'content': block_content.strip()})
                block_content = ""

    # Save the last lesson's content after the loop ends
    if current_lesson and lesson_blocks:
        current_lesson.content = json.dumps(lesson_blocks)
        current_lesson.save()

    # Update course units and hours
    course.units = subcourse_order
    course.hours = subcourse_order  # Assuming 1 hour per subcourse
    course.save()






from django.shortcuts import get_object_or_404, redirect, render
import json

def edit_lesson(request, lesson_id):
    lesson = get_object_or_404(Lesson, id=lesson_id)

    if request.method == 'POST':
        if 'word_document' in request.FILES:
            word_file = request.FILES['word_document']
            content_blocks = extract_paragraphs_and_headers_from_word(word_file)
        else:
            lesson.description = request.POST.get('description', '')

            content_blocks = []
            for idx in range(int(request.POST.get('block_count', '0'))):
                block_type = request.POST.get(f'block_type_{idx}')
                block_order = request.POST.get(f'block_order_{idx}', idx)  # Ensure order is retrieved

                if block_type == 'paragraph':
                    content = request.POST.get(f'content_{idx}', '')
                    content_blocks.append({'type': 'paragraph', 'content': content, 'order': block_order})

                elif block_type == 'header':
                    content = request.POST.get(f'content_{idx}', '')
                    content_blocks.append({'type': 'header', 'content': content, 'order': block_order})

                elif block_type == 'question':
                    question_content = request.POST.get(f'content_{idx}', '')
                    content_blocks.append({'type': 'question', 'content': question_content, 'order': block_order})

                elif block_type == 'multiple_questions':
                    multiple_question_content = request.POST.get(f'content_{idx}', '')
                    questions_list = [q.strip() for q in multiple_question_content.split(',')]
                    content_blocks.append({'type': 'multiple_questions', 'content': questions_list, 'order': block_order})

                elif block_type == 'multiple_choice':
                    question = request.POST.get(f'question_{idx}', '')
                    correct_answer = request.POST.get(f'correct_answer_{idx}', '')
                    options = request.POST.getlist(f'option_{idx}[]')  # Get multiple options
                    content_blocks.append({
                        'type': 'multiple_choice',
                        'question': question,
                        'options': options,
                        'correct_answer': correct_answer,
                        'order': block_order
                    })

            lesson.content = json.dumps(content_blocks)
            lesson.save()

            return redirect('edit_course', lesson.parent_sub_course.parent_course.id)

    else:
        try:
            content_blocks = json.loads(lesson.content)
        except json.JSONDecodeError:
            content_blocks = []

        content_blocks = sorted(content_blocks, key=lambda x: int(x.get('order', 0)))

    context = {
        'lesson': lesson,
        'content_blocks': content_blocks,
        'block_count': len(content_blocks),
    }
    return render(request, 'customadmin/edit_lesson.html', context)





from django.shortcuts import render, redirect
from myapp.forms import CourseForm

def add_course(request):
    if request.method == 'POST':
        # Handle DOCX upload
        if 'upload_docx' in request.POST:
            if 'docx_file' in request.FILES:
                docx_file = request.FILES['docx_file']
                course_form = CourseForm(request.POST, request.FILES)
                
                if course_form.is_valid():
                    # Save the course first to avoid unsaved related object error
                    course = course_form.save()  # Save the course immediately
                    
                    # Extract subcourses and lessons from DOCX file
                    extract_subcourses_lessons_from_docx(docx_file, course)
                    
                    # Now the course, subcourses, and lessons are saved, redirect to course list
                    return redirect('course_list')

        # Handle CSV upload
        elif 'upload_csv' in request.POST:
            if 'csv_file' in request.FILES:
                csv_file = request.FILES['csv_file']
                handle_csv_upload(csv_file)  # Use your existing CSV handling function
                return redirect('course_list')

    else:
        course_form = CourseForm()
    
    return render(request, 'customadmin/add_course.html', {'course_form': course_form})

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from myapp.models import Course

@csrf_exempt
def reorder_courses(request):
    if request.method == 'POST':
        order = request.POST.getlist('order[]')

        # Loop through the received course IDs and update their order field
        for idx, course_id in enumerate(order):
            course = Course.objects.get(pk=course_id)
            course.order = idx  # Update order based on the index in the list
            course.save()

        return JsonResponse({'success': True})

    return JsonResponse({'error': 'Invalid request method'}, status=400)


import csv
from myapp.models import Course, SubCourse, Lesson

def handle_csv_upload(csv_file):
    csv_reader = csv.DictReader(csv_file.read().decode('utf-8').splitlines())
    
    for row in csv_reader:
        course_title = row.get('Course Title').strip()
        subcourse_title = row.get('Subcourse Title').strip().lower()  # Normalize the subcourse title
        lesson_title = row.get('Lesson Title').strip()
        lesson_content = row.get('Lesson Content')
        lesson_order = int(row.get('Lesson Order'))  # Ensure lesson_order is an integer
        units = row.get('Units')
        hours = row.get('Hours')

        # Fetch or create the course based on the course title, units, and hours
        course, created = Course.objects.get_or_create(
            title=course_title,
            defaults={'units': units, 'hours': hours}
        )

        # Update the course units and hours if necessary
        if not created:
            if course.units != units:
                course.units = units
            if course.hours != hours:
                course.hours = hours
            course.save()

        # Normalize the subcourse title and fetch or create the subcourse
        subcourse_title_normalized = subcourse_title.strip().lower()  # Normalized subcourse title
        sub_course, created = SubCourse.objects.get_or_create(
            title=subcourse_title_normalized,  # Store the normalized title
            parent_course=course,
            defaults={'units': units, 'hours': hours}
        )

        # If the subcourse is newly created, assign the correct order
        if created:
            # Assign order based on the number of existing subcourses under the course
            sub_course.order = SubCourse.objects.filter(parent_course=course).count()
            sub_course.save()

        # If the subcourse already exists but the units/hours are different, update them
        if not created:
            if sub_course.units != units:
                sub_course.units = units
            if sub_course.hours != hours:
                sub_course.hours = hours
            sub_course.save()

        # Check if the lesson already exists under the subcourse to avoid duplication
        lesson_exists = Lesson.objects.filter(
            title=lesson_title, 
            parent_sub_course=sub_course
        ).exists()

        # Only create the lesson if it doesn't already exist
        if not lesson_exists:
            # Check if the lesson order is 1 to mark it as the first lesson
            is_first_lesson = True if lesson_order == 1 else False

            Lesson.objects.create(
                title=lesson_title, 
                content=lesson_content, 
                parent_sub_course=sub_course, 
                order=lesson_order,
                is_first_lesson=is_first_lesson  # Mark it as the first lesson if order is 1
            )

from docx import Document
from myapp.models import Course, SubCourse, Lesson

def extract_lesson_content(word_file, course_id):
    doc = Document(word_file)
    
    # Fetch the course by course_id (assuming course already exists)
    course = Course.objects.get(id=course_id)
    
    current_subcourse = None
    current_lesson = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Identify subcourse headers (e.g., "Subcourse 1: Introduction to Podcasting")
        if "Sub course" in text:
            subcourse_title = text.split(':')[1].strip()  # Extract title after "Sub course X:"
            # Fetch the existing subcourse from the database
            current_subcourse = SubCourse.objects.filter(
                title=subcourse_title,
                parent_course=course
            ).first()

        # Identify lessons under each subcourse (e.g., "Lesson 1: What is Podcasting?")
        elif "Lesson" in text and current_subcourse:
            lesson_title = text.split(':')[1].strip()  # Extract title after "Lesson X:"
            # Fetch the existing lesson under the subcourse
            current_lesson = Lesson.objects.filter(
                title=lesson_title,
                parent_sub_course=current_subcourse
            ).first()

        # Add content to the current lesson if we're within a lesson
        elif current_lesson:
            if current_lesson.content:
                current_lesson.content += f"\n{text}"  # Append content to existing content
            else:
                current_lesson.content = text  # Add new content
            current_lesson.save()

    return "Lesson contents have been successfully updated."



def generate_course_html(course):
    context = {'course': course}
    html_content = render_to_string('myapp/course_templates/base_course_template.html', context)
    file_path = os.path.join('myapp/templates/myapp/generated_courses', f'{course.title}.html')
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    with open(file_path, 'w') as f:
        f.write(html_content)

from django.shortcuts import get_object_or_404, redirect

from django.shortcuts import redirect, get_object_or_404

def add_sub_course(request):
    # Fetch the course using course_id from GET parameters
    course_id = request.GET.get('course_id')
    parent_course = get_object_or_404(Course, id=course_id)
    
    if request.method == 'POST':
        form = SubCourseForm(request.POST)
        if form.is_valid():
            sub_course = form.save(commit=False)
            sub_course.parent_course = parent_course  # Link the new sub-course to the parent course
            sub_course.order = SubCourse.objects.filter(parent_course=parent_course).count() + 1
            sub_course.save()
            return redirect('edit_course', course_id=parent_course.id)
    else:
        form = SubCourseForm()

    return render(request, 'customadmin/add_sub_course.html', {'form': form, 'course': parent_course})


def edit_sub_course(request, sub_course_id):
    sub_course = get_object_or_404(SubCourse, id=sub_course_id)
    if request.method == 'POST':
        form = SubCourseForm(request.POST, instance=sub_course)
        if form.is_valid():
            form.save()
            return redirect('edit_course', course_id=sub_course.parent_course.id)
    else:
        form = SubCourseForm(instance=sub_course)
    return render(request, 'customadmin/edit_sub_course.html', {'form': form, 'sub_course': sub_course})

import os
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings

def handle_uploaded_file(file):
    upload_dir = '/Users/Julia/Downloads/braine-package/myapp/static/myapp/images/lesson_images/'
    
    # Ensure the directory exists
    if not os.path.exists(upload_dir):
        os.makedirs(upload_dir)
    
    file_path = os.path.join(upload_dir, file.name)
    
    # Save the file in the specified directory
    with open(file_path, 'wb+') as destination:
        for chunk in file.chunks():
            destination.write(chunk)
    
    # Return the relative file path for template use
    return f'/static/myapp/images/lesson_images/{file.name}'

def add_lesson(request):
    sub_course_id = request.GET.get('sub_course_id')
    sub_course = get_object_or_404(SubCourse, id=sub_course_id)
    
    if request.method == 'POST':
        form = LessonForm(request.POST)
        if form.is_valid():
            lesson = form.save(commit=False)
            lesson.parent_sub_course = sub_course
            # Assign the next available order for the lesson under the sub-course
            lesson.order = Lesson.objects.filter(parent_sub_course=sub_course).count() + 1
            lesson.save()
            return redirect('edit_course', course_id=sub_course.parent_course.id)  # Redirect to course detail page
    else:
        form = LessonForm()

    return render(request, 'customadmin/add_lesson.html', {'form': form, 'sub_course': sub_course})


import json
from docx import Document
from django.shortcuts import get_object_or_404, redirect, render

def extract_paragraphs_and_headers_from_word(word_file):
    doc = Document(word_file)
    content_blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if any part of the paragraph is bold
        is_bold = any(run.bold for run in para.runs if run.bold is not None)

        # Classify as header or paragraph
        if is_bold:
            block_type = 'header'
        else:
            block_type = 'paragraph'
        
        content_blocks.append({'type': block_type, 'content': text})

    return content_blocks


import re

def convert_to_embed_url(url):
    youtube_regex = r'(https?://)?(www\.)?(youtube|youtu|youtube-nocookie)\.(com|be)/(watch\?v=|embed/|v/|.+\?v=)?([^&=%\?]{11})'
    match = re.match(youtube_regex, url)
    if match:
        video_id = match.group(6)
        # Append parameters to remove branding and controls
        return f"https://www.youtube.com/embed/{video_id}?rel=0&modestbranding=1&controls=0&fs=0&disablekb=1&autohide=1"
    return url



from django.shortcuts import get_object_or_404, redirect, render
import json

def edit_lesson(request, lesson_id):
    lesson = get_object_or_404(Lesson, id=lesson_id)

    if request.method == 'POST':
        # Check if a Word document is uploaded (existing logic)
        if 'word_document' in request.FILES:
            word_file = request.FILES['word_document']
            content_blocks = extract_paragraphs_and_headers_from_word(word_file)
        else:
            # Regular form submission to handle the lesson content
            lesson.description = request.POST.get('description', '')

            content_blocks = []
            for idx in range(int(request.POST.get('block_count', '0'))):
                block_type = request.POST.get(f'block_type_{idx}')
                block_order = request.POST.get(f'block_order_{idx}', idx)  # Ensure order is retrieved

                if block_type == 'paragraph':
                    content = request.POST.get(f'content_{idx}', '')
                    content_blocks.append({'type': 'paragraph', 'content': content, 'order': block_order})

                elif block_type == 'image':
                    image = request.FILES.get(f'image_{idx}')
                    if image:
                        image_url = handle_uploaded_file(image)
                        content_blocks.append({'type': 'image', 'content': image_url, 'order': block_order})

                elif block_type == 'header':
                    content = request.POST.get(f'content_{idx}', '')
                    content_blocks.append({'type': 'header', 'content': content, 'order': block_order})

                elif block_type == 'task':
                    task_content = request.POST.get(f'content_{idx}', '')
                    content_blocks.append({'type': 'task', 'content': task_content, 'order': block_order})

                elif block_type == 'question':
                    question_content = request.POST.get(f'content_{idx}', '')
                    content_blocks.append({'type': 'question', 'content': question_content, 'order': block_order})

                elif block_type == 'multiple_questions':
                    multiple_question_content = request.POST.get(f'content_{idx}', '')
                    questions_list = [q.strip() for q in multiple_question_content.split(',')]
                    content_blocks.append({'type': 'multiple_questions', 'content': questions_list, 'order': block_order})

                elif block_type == 'multiple_choice':
                    question = request.POST.get(f'question_{idx}', '')
                    correct_answer = request.POST.get(f'correct_answer_{idx}', '')
                    options = request.POST.getlist(f'option_{idx}[]')
                    content_blocks.append({
                        'type': 'multiple_choice',
                        'question': question,
                        'options': options,
                        'correct_answer': correct_answer,
                        'order': block_order
                    })

                # Handle video URL
                elif block_type == 'video':
                    video_url = request.POST.get(f'video_url_{idx}', '')
                    content_blocks.append({'type': 'video', 'content': video_url, 'order': block_order})

                # Handle Reflection section
                elif "Reflection" in block_type:
                    reflection_content = request.POST.get(f'content_{idx}', '')
                    reflection_content = reflection_content.replace("End of Sub course", f"End of {lesson.title}")
                    content_blocks.append({'type': 'reflection', 'content': reflection_content, 'order': block_order})

                # Handle Course Wrap-Up or Congratulations
                elif "Course Wrap-Up" in block_type or "Congratulations" in block_type:
                    wrap_up_content = request.POST.get(f'content_{idx}', '')
                    content_blocks.append({'type': 'course_wrap_up', 'content': wrap_up_content, 'order': block_order})

            # Save the updated content blocks in JSON format with order
            lesson.content = json.dumps(content_blocks)
            lesson.save()

            # Redirect back to the course detail page after saving
            return redirect('edit_course', lesson.parent_sub_course.parent_course.id)

    else:
        # Load existing content blocks if available
        try:
            content_blocks = json.loads(lesson.content)
        except json.JSONDecodeError:
            content_blocks = []

        # Sort the content blocks by the 'order' field before passing them to the template
        content_blocks = sorted(content_blocks, key=lambda x: int(x.get('order', 0)))

    context = {
        'lesson': lesson,
        'content_blocks': content_blocks,
        'block_count': len(content_blocks),  # Pass the number of blocks to the template
    }
    return render(request, 'customadmin/edit_lesson.html', context)


from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from myapp.models import UserAnswer

@csrf_exempt
def process_answer(request, lesson_id, question_id):
    lesson = get_object_or_404(Lesson, id=lesson_id)
    content_blocks = json.loads(lesson.content)
    
    # Ensure question_id is valid
    if question_id >= len(content_blocks) or question_id < 0:
        return JsonResponse({'error': 'Invalid question ID'}, status=404)
    
    question_block = content_blocks[question_id]

    # Ensure it's a multiple-choice question
    if question_block.get('type') != 'multiple_choice':
        return JsonResponse({'error': 'Question is not a multiple-choice type'}, status=400)

    user_answer = request.POST.get(f'answer_{question_id}', '').strip().lower()
    correct_answer = question_block.get('correct_answer', '').strip().lower()

    is_correct = user_answer == correct_answer

    # Save the user's answer to the database
    UserAnswer.objects.create(
        user=request.user,
        lesson=lesson,
        question_type=question_block.get('type'),
        question_content=question_block.get('question'),
        user_answer=user_answer,
        correct_answer=correct_answer,
        is_correct=is_correct
    )

    # Return the response as JSON
    return JsonResponse({
        'correct': is_correct,
        'correct_answer': correct_answer
    })

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from myapp.models import Lesson, ContentBlock

@csrf_exempt
def save_block(request, block_id):
    if request.method == 'POST':
        block_type = request.POST.get('block_type')
        content = request.POST.get('content')
        question = request.POST.get('question')
        correct_answer = request.POST.get('correct_answer')

        # Retrieve or create the block
        block = ContentBlock.objects.get(id=block_id)

        # Update the block data
        block.type = block_type
        block.content = content
        block.question = question
        block.correct_answer = correct_answer
        block.save()

        return JsonResponse({'success': True})

    return JsonResponse({'success': False, 'error': 'Invalid request'})

from django.shortcuts import get_object_or_404, redirect
from django.http import HttpResponseRedirect

def update_sub_course_order(request, sub_course_id):
    sub_course = get_object_or_404(SubCourse, id=sub_course_id)
    if request.method == 'POST':
        new_order = request.POST.get('order')
        if new_order:
            sub_course.order = int(new_order)
            sub_course.save()
    return redirect('edit_course', course_id=sub_course.parent_course.id)

def update_lesson_order(request, lesson_id):
    lesson = get_object_or_404(Lesson, id=lesson_id)
    if request.method == 'POST':
        new_order = request.POST.get('order')
        if new_order:
            lesson.order = int(new_order)
            lesson.save()
    return redirect('edit_course', course_id=lesson.parent_sub_course.parent_course.id)


def edit_course(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    if request.method == 'POST':
        form = CourseForm(request.POST, request.FILES, instance=course)
        if form.is_valid():
            form.save()
            return redirect('edit_course', course_id=course.id)
    else:
        form = CourseForm(instance=course)
    sub_courses = course.sub_courses.all()
    return render(request, 'customadmin/edit_course.html', {'course': course, 'form': form, 'sub_courses': sub_courses})

def delete_course(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    if request.method == 'POST':
        # Delete the generated HTML file
        file_path = os.path.join(settings.BASE_DIR, 'myapp/templates/myapp/generated_courses', f'{course.title}.html')
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Delete the course
        course.delete()
        return redirect('dashboard')
    return render(request, 'customadmin/confirm_delete.html', {'course': course})

class CustomLoginView(LoginView):
    template_name = 'customadmin/login.html'

class CustomLogoutView(LogoutView):
    next_page = reverse_lazy('dashboard')

# customadmin/views.py
import os
from django.conf import settings

def handle_uploaded_file(f):
    # Define the directory path relative to MEDIA_ROOT
    directory = os.path.join(settings.MEDIA_ROOT, 'lesson_images')
    
    # Create the directory if it doesn't exist
    if not os.path.exists(directory):
        os.makedirs(directory)
    
    # Define the full file path
    file_path = os.path.join(directory, f.name)
    
    # Save the file to the appropriate location
    with open(file_path, 'wb+') as destination:
        for chunk in f.chunks():
            destination.write(chunk)
    
    # Return the relative URL path
    return os.path.join(settings.MEDIA_URL, 'lesson_images', f.name)


from django.shortcuts import render, redirect
from django.contrib import messages
from myapp.forms import CSVUploadForm
from myapp.models import Course, SubCourse, Lesson
import csv


def upload_csv(request):
    if request.method == 'POST':
        form = CSVUploadForm(request.POST, request.FILES)
        if form.is_valid():
            csv_file = request.FILES['csv_file']
            if not csv_file.name.endswith('.csv'):
                messages.error(request, 'The file is not CSV type')
                return redirect('upload_csv')
            try:
                data = csv_file.read().decode('utf-8')
                reader = csv.DictReader(data.splitlines())
                for row in reader:
                    if 'course' in request.path:
                        Course.objects.create(
                            title=row['title'],
                            description=row['description'],
                            units=row['units'],
                            hours=row['hours']
                        )
                    elif 'sub_course' in request.path:
                        SubCourse.objects.create(
                            parent_course=Course.objects.get(title=row['parent_course']),
                            title=row['title'],
                            description=row['description'],
                            units=row['units'],
                            hours=row['hours']
                        )
                    elif 'lesson' in request.path:
                        Lesson.objects.create(
                            parent_sub_course=SubCourse.objects.get(title=row['parent_sub_course']),
                            title=row['title'],
                            content=row['content']
                        )
                messages.success(request, 'CSV file processed successfully')
                return redirect('some_view')
            except Exception as e:
                messages.error(request, f'Error processing CSV file: {e}')
                return redirect('upload_csv')
    else:
        form = CSVUploadForm()

    return render(request, 'upload_csv.html', {'form': form})


from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.models import User
from django.core.mail import send_mail
from myapp.models import EmailCollection

def user_management(request):
    users = User.objects.all()
    email_collections = EmailCollection.objects.all()
    context = {
        'users': users,
        'email_collections': email_collections,
    }
    return render(request, 'customadmin/user_management.html', context)

def send_password_reset(request, user_id):
    user = get_object_or_404(User, id=user_id)
    subject = 'Password Reset Request'
    message = f'Click the link below to reset your password:\n\n{request.build_absolute_uri(reverse("password_reset_confirm", args=[user.id]))}'
    send_mail(subject, message, 'from@example.com', [user.email])
    return redirect('user_management')


# customadmin/views.py


from django.shortcuts import render, get_object_or_404
from django.contrib.auth.models import User
from myapp.models import EmailCollection, QuizResponse  # Import the QuizResponse model

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.models import User
from myapp.models import QuizResponse

def view_user_quiz_details(request, user_id):
    user = get_object_or_404(User, id=user_id)
    
    try:
        quiz_response = QuizResponse.objects.get(user=user)
    except QuizResponse.DoesNotExist:
        quiz_response = None  # Handle the case where there's no QuizResponse

    return render(request, 'customadmin/view_user_quiz_details.html', {
        'user': user,
        'quiz_response': quiz_response,
    })


def edit_user(request, user_id):
    user = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        # Handle user update logic here
        user.email = request.POST.get('email', user.email)
        user.save()
        return redirect('customadmin_user_management')
    return render(request, 'customadmin/edit_user.html', {'user': user})

from django.shortcuts import get_object_or_404, redirect, render
from django.contrib.auth.models import User  # Assuming User is the model you're deleting

def delete_user(request, item_id):
    item = get_object_or_404(User, id=item_id)  # Adjust the model as needed
    if request.method == 'POST':
        item.delete()
        return redirect('user_management')  # Replace with the correct redirect URL
    return render(request, 'customadmin/delete_user.html', {'item': item})


from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.models import User
from django.http import HttpResponseBadRequest

def delete_multiple_users(request):
    if request.method == 'POST':
        user_ids = request.POST.getlist('user_ids')
        if not user_ids:
            return HttpResponseBadRequest('No users selected')

        User.objects.filter(id__in=user_ids).delete()
        return redirect('user_management')
    return HttpResponseBadRequest('Invalid request method')


from django.shortcuts import render, get_object_or_404, redirect
from myapp.models import KnowledgeBaseCategory, KnowledgeBaseSubCategory, KnowledgeBaseArticle
from myapp.forms import KnowledgeBaseCategoryForm, KnowledgeBaseSubCategoryForm, KnowledgeBaseArticleForm

# Manage Knowledge Base View
def manage_knowledge_base(request):
    categories = KnowledgeBaseCategory.objects.all()
    subcategories = KnowledgeBaseSubCategory.objects.all()
    articles = KnowledgeBaseArticle.objects.all()

    # Debugging prints
    print("Categories:", categories)
    print("Subcategories:", subcategories)
    print("Articles:", articles)

    context = {
        'categories': categories,
        'subcategories': subcategories,
        'articles': articles,
    }
    return render(request, 'customadmin/manage_knowledge_base.html', context)

# Add Category View
def add_category(request):
    if request.method == 'POST':
        form = KnowledgeBaseCategoryForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('manage_knowledge_base')
    else:
        form = KnowledgeBaseCategoryForm()
    return render(request, 'customadmin/add_category.html', {'form': form})

# Edit Category View
def edit_category(request, id=None):
    category = None
    if id:
        category = get_object_or_404(KnowledgeBaseCategory, id=id)

    if request.method == 'POST':
        form = KnowledgeBaseCategoryForm(request.POST, request.FILES, instance=category)
        if form.is_valid():
            form.save()
            return redirect('manage_knowledge_base')
    else:
        form = KnowledgeBaseCategoryForm(instance=category)

    return render(request, 'customadmin/edit_category.html', {
        'form': form,
        'category': category,
    })

# Add Subcategory View
def add_subcategory(request):
    categories = KnowledgeBaseCategory.objects.all()

    if request.method == 'POST':
        category_id = request.POST['category']
        title = request.POST['title']
        description = request.POST['description']
        icon = request.FILES.get('icon')

        category = get_object_or_404(KnowledgeBaseCategory, id=category_id)
        KnowledgeBaseSubCategory.objects.create(
            category=category,
            title=title,
            description=description,
            icon=icon
        )
        return redirect('manage_knowledge_base')

    return render(request, 'customadmin/add_subcategory.html', {'categories': categories})

# Edit Subcategory View
def edit_subcategory(request, id):
    subcategory = get_object_or_404(KnowledgeBaseSubCategory, id=id)
    if request.method == 'POST':
        form = KnowledgeBaseSubCategoryForm(request.POST, instance=subcategory)
        if form.is_valid():
            form.save()
            return redirect('manage_knowledge_base')
    else:
        form = KnowledgeBaseSubCategoryForm(instance=subcategory)
    return render(request, 'customadmin/edit_subcategory.html', {'form': form})

# Add Article View
def add_article(request):
    subcategories = KnowledgeBaseSubCategory.objects.all()
    if request.method == 'POST':
        form = KnowledgeBaseArticleForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('manage_knowledge_base')
    else:
        form = KnowledgeBaseArticleForm()
    return render(request, 'customadmin/add_article.html', {'form': form, 'subcategories': subcategories})

# Edit Article View

from django.shortcuts import render, get_object_or_404, redirect
from myapp.models import KnowledgeBaseArticle
from myapp.forms import KnowledgeBaseArticleForm
import json

def edit_article(request, id):
    article = get_object_or_404(KnowledgeBaseArticle, id=id)
    
    if request.method == 'POST':
        form = KnowledgeBaseArticleForm(request.POST, instance=article)
        if form.is_valid():
            # Save the basic form data
            article = form.save(commit=False)
            
            # Process the content blocks
            block_count = int(request.POST.get('block_count', 0))
            content_blocks = []
            for i in range(block_count):
                block_type = request.POST.get(f'block_type_{i}')
                if block_type == 'paragraph':
                    content = request.POST.get(f'content_{i}', '')
                    content_blocks.append({'type': 'paragraph', 'content': content})
                elif block_type == 'image':
                    image = request.FILES.get(f'image_{i}')
                    if image:
                        # Handle image upload and get the URL
                        image_url = handle_uploaded_file(image)  # Assuming this is defined elsewhere
                        content_blocks.append({'type': 'image', 'content': image_url})
                elif block_type == 'header':
                    content = request.POST.get(f'content_{i}', '')
                    content_blocks.append({'type': 'header', 'content': content})
            
            # Save the content blocks as JSON in the article's content field
            article.content = json.dumps(content_blocks)
            article.save()

            return redirect('manage_knowledge_base')
    else:
        form = KnowledgeBaseArticleForm(instance=article)

    try:
        content_blocks = json.loads(article.content)  # Assuming content is stored as JSON
    except json.JSONDecodeError:
        content_blocks = []

    context = {
        'form': form,
        'article': article,
        'content_blocks': content_blocks,
    }

    return render(request, 'customadmin/edit_article.html', context)


from django.shortcuts import redirect, get_object_or_404
from myapp.models import KnowledgeBaseCategory, KnowledgeBaseSubCategory, KnowledgeBaseArticle


def delete_category(request, category_id):
    category = get_object_or_404(KnowledgeBaseCategory, id=category_id)
    if request.method == 'POST':
        category.delete()
    return redirect('manage_knowledge_base')

def delete_subcategory(request, id):
    subcategory = get_object_or_404(KnowledgeBaseSubCategory, id=id)
    if request.method == 'POST':
        subcategory.delete()
    return redirect('manage_knowledge_base')

def delete_article(request, article_id):
    article = get_object_or_404(KnowledgeBaseArticle, id=article_id)
    if request.method == 'POST':
        article.delete()
    return redirect('manage_knowledge_base')


import csv
from django.http import HttpResponse, JsonResponse
from myapp.models import Transaction

# View for downloading transactions as a CSV file
def download_transactions_csv(request):
    # Create the HttpResponse object with the appropriate CSV header.
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="transactions.csv"'

    # Create a CSV writer
    writer = csv.writer(response)

    # Write the header row
    writer.writerow(['User', 'Subscription Type', 'Amount', 'Status', 'Transaction Date', 'Recurring', 'Next Billing Date', 'Error Logs'])

    # Write data rows
    transactions = Transaction.objects.all().values_list(
        'user__email', 'subscription_type', 'amount', 'status', 'created_at', 'recurring', 'next_billing_date', 'error_logs'
    )
    for transaction in transactions:
        writer.writerow(transaction)

    return response


# API view to get transactions data for rendering in HTML
def view_transactions(request):
    status = request.GET.get('status', None)
    if status:
        transactions = Transaction.objects.filter(status=status)
    else:
        transactions = Transaction.objects.all()

    transactions_list = [
        {
            "user": transaction.user.email,
            "subscription_type": transaction.subscription_type,
            "amount": transaction.amount,
            "status": transaction.status,
            "transaction_date": transaction.created_at.strftime("%Y-%m-%d %H:%M:%S"),
            "recurring": transaction.recurring,
            "next_billing_date": transaction.next_billing_date.strftime("%Y-%m-%d") if transaction.next_billing_date else None,
            "error_logs": transaction.error_logs,
        }
        for transaction in transactions
    ]

    return JsonResponse({"transactions": transactions_list})


# views.py
# views.py
from django.shortcuts import render
from myapp.models import Transaction

def customadmin_transactions(request):
    # Fetch all transactions from the database
    transactions = Transaction.objects.all()
    
    return render(request, 'customadmin/transactions.html', {'transactions': transactions})


def delete_sub_course(request, sub_course_id):
    sub_course = get_object_or_404(SubCourse, id=sub_course_id)
    if request.method == 'POST':
        sub_course.delete()
        return redirect('edit_course', course_id=sub_course.parent_course.id)
    return render(request, 'customadmin/confirm_delete.html', {'sub_course': sub_course})

def delete_lesson(request, lesson_id):
    lesson = get_object_or_404(Lesson, id=lesson_id)
    if request.method == 'POST':
        lesson.delete()
        return redirect('edit_course', course_id=lesson.parent_sub_course.parent_course.id)
    return render(request, 'customadmin/confirm_delete.html', {'lesson': lesson})


# views.py
from django.shortcuts import render, redirect
from myapp.models import BlogPost  # Make sure to import your BlogPost model

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from myapp.models import BlogPost  # Import your BlogPost model
from django.contrib import messages  # For displaying messages

@login_required
def add_blog(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        content = []  # Initialize content as a list

        # Handle content blocks
        block_types = request.POST.getlist('block_type')
        block_contents = request.POST.getlist('content')

        for block_type, block_content in zip(block_types, block_contents):
            content.append({
                'type': block_type,
                'content': block_content,
            })

        print("Content being saved:", content)  # Log the content

        # Create the blog post
        BlogPost.objects.create(
            title=title,
            author=request.user,  # Assign the currently logged-in user as the author
            content=content
        )
        messages.success(request, 'Blog post added successfully!')
        return redirect('blog_list')  # Redirect to the list of blog posts after saving

    return render(request, 'customadmin/add_blog.html')


# views.py
from django.shortcuts import render, get_object_or_404, redirect
from myapp.models import BlogPost  # Ensure you have a BlogPost model
from django.contrib import messages  # For displaying messages

def blog_list(request):
    blogs = BlogPost.objects.all()  # Get all blog posts
    return render(request, 'customadmin/blog_list.html', {'blogs': blogs})


from django.shortcuts import get_object_or_404, redirect, render
from myapp.models import BlogPost  # Ensure your BlogPost model is imported
from django.contrib import messages  # For displaying messages
import json  # Import for handling JSON data

@login_required
def edit_blog(request, blog_id):
    blog_post = get_object_or_404(BlogPost, id=blog_id)

    if request.method == 'POST':
        blog_post.title = request.POST.get('title')
        
        # Check if the content is a JSON string and parse it, otherwise assume it's already a list
        content = request.POST.getlist('content')  # Assuming this is how you're sending the content from the form
        block_types = request.POST.getlist('block_type')

        content_data = []  # Initialize content as a list
        for block_type, block_content in zip(block_types, content):
            content_data.append({
                'type': block_type,
                'content': block_content,
            })

        # Save the updated content as a JSON string
        blog_post.content = json.dumps(content_data)
        
        # Fetch the current user instance and assign it to the blog post
        blog_post.author = request.user  # Assuming you want to set the logged-in user as the author

        try:
            blog_post.save()
            messages.success(request, 'Blog post updated successfully!')
            return redirect('blog_detail', post_id=blog_post.id)  # Adjust this as needed
        except Exception as e:
            messages.error(request, f"Error updating blog post: {e}")

    else:
        # Load existing content blocks, ensuring it's handled properly
        if isinstance(blog_post.content, str):  # If it's a string, decode it
            content_blocks = json.loads(blog_post.content)
        else:
            content_blocks = blog_post.content  # Assuming it's already a list

    return render(request, 'customadmin/edit_blog.html', {'blog_post': blog_post, 'content_blocks': content_blocks})


def delete_blog(request, blog_id):
    blog = get_object_or_404(BlogPost, id=blog_id)

    if request.method == 'POST':
        blog.delete()
        messages.success(request, 'Blog deleted successfully!')
        return redirect('blog_list')

    return render(request, 'customadmin/delete_blog.html', {'blog': blog})
